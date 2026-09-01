"""Markdown is the canonical source for generated Word documentation."""

from pathlib import Path

from docx import Document

from scripts.export_docx import export


def test_export_docx_renders_markdown(tmp_path):
    source = tmp_path / "guide.md"
    source.write_text(
        "# Guide\n\nA **single** source with `code`.\n\n"
        "| Control | Status |\n|---|---|\n| DPA | Missing |\n",
        encoding="utf-8",
    )
    outputs = export([source], tmp_path / "out")
    assert outputs == [tmp_path / "out" / "guide.docx"]
    document = Document(outputs[0])
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Guide" in text
    assert "single source with code" in text
    assert document.tables[0].cell(1, 0).text == "DPA"


def test_legacy_duplicate_generators_are_removed():
    docs = Path(__file__).resolve().parents[1] / "docs"
    assert not list(docs.glob("build_*.py"))
