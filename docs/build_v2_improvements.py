"""Generate the v2 improvements Word doc: persistence, DPIA Workspace, Evidence Repository."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "V2_Improvements_AI_Governance.docx"
NAVY = RGBColor(0x1B, 0x21, 0x18)
ACCENT = RGBColor(0x3D, 0x5A, 0x1F)
MUTED = RGBColor(0x4A, 0x52, 0x45)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1B2118"
ROW_ALT = "F3F5EE"


def set_run(run, size=11, bold=False, color=NAVY, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_para(
    doc,
    text,
    size=11,
    bold=False,
    color=NAVY,
    italic=False,
    space_after=8,
    space_before=0,
    align=None,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"})
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=NAVY, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, color=WHITE, fill=HEADER_BG)
    for r_idx, row in enumerate(rows):
        fill = ROW_ALT if r_idx % 2 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=10, fill=fill)
    doc.add_paragraph()


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT if level == 1 else NAVY
        run.font.name = "Calibri"
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True, size=11)
        r2 = p.add_run(text)
        set_run(r2, size=11)
    else:
        r = p.add_run(text)
        set_run(r, size=11)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

    add_para(
        doc,
        "RESPONSIBLE AI GOVERNANCE  ·  V2 ROADMAP",
        size=11,
        bold=True,
        color=ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "V2 improvements",
        size=26,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "From a production-oriented reference implementation to an enterprise control plane",
        size=13,
        italic=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16,
    )
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Date", date.today().isoformat()],
            ["Audience", "Reviewers, hiring managers, and anyone scoring this as a portfolio product"],
            ["Honest label today", "Production-oriented reference implementation (Demo / PoC). Not a production-ready enterprise platform."],
            ["Target after v2", "Durable, tenant-isolated control plane with an electronic DPIA workspace and a real evidence repository"],
            ["What must not change", "Evidence → controls → risk factors → deterministic score → triage → human Jira gates. LLM never decides."],
        ],
    )

    heading(doc, "1. Why this document exists", 1)
    add_para(
        doc,
        "v1 already has the hard parts that most “AI compliance” demos skip: a deterministic engine, "
        "EvidenceItem distinct from findings, PrivacyTriage as a first-class object, DecisionRecord, "
        "Jira department gates, and EVIDENCE NOT FOUND instead of invented certifications.",
    )
    add_para(
        doc,
        "That is why the product scores well as architecture. It still loses points in three places that "
        "an auditor, a DPO, or a sceptical reviewer will find in minutes. This document specifies those "
        "gaps without rewriting the product story. Closing them is the difference between a strong "
        "reference implementation (~8.5) and a credible enterprise control plane (~9.0+).",
    )
    add_table(
        doc,
        ["#", "Gap", "What v1 actually does", "Why it still loses points"],
        [
            [
                "1",
                "Persistence",
                "In-memory store keyed by assessment_id, with TTL and an access token.",
                "Restart, second worker, or tenant isolation still fail. Durable audit is not possible in RAM.",
            ],
            [
                "2",
                "Electronic DPIA",
                "CTRL-DPIA checks status + reference. PrivacyTriage can say “DPIA review indicated”.",
                "That is a control over a DPIA, not a DPIA. No screening, risks, residual risk, DPO approval, or version.",
            ],
            [
                "3",
                "Documentary evidence",
                "EvidenceItem records control_id, document name, excerpt, status, timestamp, confidence.",
                "There is no file. SOC 2, DPA PDF, ISO, model cards, and clauses stay missing or declared_unverified.",
            ],
        ],
    )

    heading(doc, "2. Commercial framing — keep this honest", 1)
    add_para(doc, "Do not sell v1 as:", bold=True, space_after=4)
    add_para(doc, "“Production-ready enterprise platform.”", italic=True)
    add_para(doc, "Sell v1 as:", bold=True, space_after=4, space_before=8)
    add_para(
        doc,
        "Production-oriented reference implementation of an Enterprise AI Governance & Risk Control Layer.",
        italic=True,
    )
    add_para(
        doc,
        "The architecture, the ADRs, the conservative scoring, and the Jira gates are production-shaped. "
        "The runtime is not. Reviewers respect that distinction. Inflating the claim is the fastest way "
        "to lose the points the design has already earned.",
        space_before=8,
    )
    add_table(
        doc,
        ["Claim", "v1", "After v2"],
        [
            ["Reproducible triage", "Yes, for one process", "Yes, across workers and restarts"],
            ["Multi-tenant isolation", "No", "PostgreSQL + tenant_id + RLS"],
            ["DPIA as a work product", "Status / reference only", "Full workspace with version and DPO sign-off"],
            ["Evidence as a file", "No uploads", "Hashed artifacts bound to controls"],
            ["Audit question: show the pack", "JSON evidence_pack", "JSON + files + SHA-256 + reviewer decision"],
            ["Engine APPROVE", "Never", "Still never"],
        ],
    )

    heading(doc, "3. Gap 1 — Durable persistence and tenant isolation", 1)
    heading(doc, "Current state", 2)
    add_para(
        doc,
        "app/store.py is no longer a naive singleton. It keys sessions by assessment_id, uses a threading "
        "lock, expires sessions, caps memory, and can require an access token. That is enough for a demo "
        "and for concurrent local tests. It is still RAM. Production must replace this module.",
    )
    bullet(doc, "Restarting uvicorn erases assessments, DecisionRecords, and webhook event ids.")
    bullet(doc, "Two uvicorn workers do not share memory; the same assessment_id is not one object.")
    bullet(doc, "There is no tenant_id. One process is one implicit tenant.")
    bullet(doc, "DecisionRecord can be overwritten in place when a Jira webhook arrives. That is not an append-only audit log.")
    bullet(doc, "Chat history is client-side. Correct for privacy in a PoC; insufficient if Legal later asks “what did the assistant say on Tuesday?”")

    heading(doc, "v2 target", 2)
    add_para(
        doc,
        "Replace the store, not the scoring engine. Pydantic models stay the governance contract (ADR-7). "
        "PostgreSQL is the source of truth. Redis, if used, is only a UI cache.",
    )
    add_table(
        doc,
        ["Object", "Table / store", "Rules"],
        [
            ["Tenant", "tenants", "Stable id, name, region, retention policy. All other rows carry tenant_id."],
            ["Vendor", "vendors", "Unique per tenant. Re-assessments attach to the same vendor, not a new orphan."],
            ["Assessment", "assessments", "UUID PK. intake JSON + engine_version + created_at. Never mutate intake after score."],
            ["DecisionRecord", "decision_records", "INSERT only. Correction = new row with supersedes_id. DB trigger blocks UPDATE/DELETE."],
            ["Jira mapping", "jira_issues", "issue.key → assessment_id. Webhooks must not bind to “latest in RAM”."],
            ["Webhook events", "processed_events", "event_id unique per tenant. Survives restart (today it does not)."],
            ["Audit hash", "evidence_packs", "Canonical JSON + sha256. Export includes engine version."],
        ],
    )
    heading(doc, "Tenant isolation (non-negotiable)", 2)
    bullet(doc, " every query is scoped. No SELECT without tenant_id in the WHERE clause.", bold_prefix="Row-level security:")
    bullet(doc, " access tokens become short-lived session credentials, not the isolation boundary.", bold_prefix="Auth:")
    bullet(doc, " SSO (OIDC) maps user → tenant + RACI role (assessor, legal, infosec, aigov, auditor, admin).", bold_prefix="Identity:")
    bullet(doc, " demo tenants and live tenants must not share a database without schema or RLS separation.", bold_prefix="Environments:")
    heading(doc, "Acceptance criteria", 2)
    bullet(doc, "Restarting the API does not lose assessments, packs, or processed webhook ids.")
    bullet(doc, "Two workers return the same assessment for the same assessment_id.")
    bullet(doc, "Tenant A cannot read Tenant B’s assessments by guessing a UUID.")
    bullet(doc, "Direct UPDATE on decision_records fails or emits a tamper event.")
    bullet(doc, "Jira webhook resolves assessment_id from issue.key, never from “latest session”.")

    heading(doc, "4. Gap 2 — DPIA Workspace / DPIA Generator", 1)
    heading(doc, "Current state — a control over the DPIA, not a DPIA", 2)
    add_para(
        doc,
        "ADR-10 was the right call: GDPR is not a has_dpa checkbox. PrivacyTriage already captures purposes, "
        "subjects, Art. 9, transfers, retention, and automated decisions. CTRL-DPIA can fail the vendor when "
        "a DPIA is indicated and no reference exists.",
    )
    add_para(
        doc,
        "That is still a gate, not a work product. The engine can say “DPIA review indicated” and check "
        "dpia_status / dpia_reference. It does not generate, version, or manage an electronic DPIA. "
        "A DPO cannot complete Art. 35 work inside this console. Legal’s Jira ticket still says "
        "“Review DPA and approve DPIA” while the DPIA itself lives in email or a Word file elsewhere.",
    )
    add_table(
        doc,
        ["v1 object", "What it is", "What it is not"],
        [
            ["PrivacyTriage", "Screening signal used by the engine", "A DPIA document"],
            ["CTRL-DPIA", "Control: is a completed DPIA evidenced?", "The assessment of necessity, risks, and residual risk"],
            ["dpia_reference", "External id string", "Versioned content with DPO sign-off"],
            ["privacy_risk_level", "Business-declared label; does not override score", "DPO residual-risk determination"],
        ],
    )

    heading(doc, "v2 target — electronic DPIA lifecycle", 2)
    add_para(
        doc,
        "Add a DPIA Workspace that is triggered by PrivacyTriage and becomes the evidence for CTRL-DPIA. "
        "The engine still does not approve the vendor. The DPO approves the DPIA. The engine only checks "
        "that an approved, in-date DPIA exists when one is indicated.",
    )
    add_table(
        doc,
        ["Stage", "Owner", "Required output", "Engine effect"],
        [
            [
                "1. Trigger",
                "Engine",
                "privacy_assessment_required + dpia_decision_basis from current intake",
                "Opens workspace when indicated; never skips Legal",
            ],
            [
                "2. Screening",
                "Assessor + Legal",
                "Art. 35 / WP248 criteria: systematic evaluation, special categories, large scale, vulnerable subjects, ADM",
                "If screening = not required, CTRL-DPIA may become not_applicable after DPO confirms",
            ],
            [
                "3. Necessity & proportionality",
                "Business + Legal",
                "Purposes, lawful basis, data minimisation, alternatives considered",
                "Incomplete stage keeps CTRL-DPIA = fail",
            ],
            [
                "4. Risks to data subjects",
                "Assessor + DPO",
                "Risk register: likelihood × severity for identified harms (discrimination, re-identification, chilling effect, etc.)",
                "No risks documented → DPIA cannot be marked complete",
            ],
            [
                "5. Controls",
                "Legal + InfoSec",
                "Each risk mapped to a control and, where possible, to an EvidenceArtifact",
                "Unbound controls stay residual; they do not silently pass",
            ],
            [
                "6. Residual risk",
                "DPO",
                "DPO residual-risk statement. Distinct from engine residual risk.",
                "Engine score is unchanged by this field. Both are stored.",
            ],
            [
                "7. DPO review",
                "DPO",
                "Review notes, consultation record, objections",
                "CTRL-DPIA stays fail until this stage is closed",
            ],
            [
                "8. Approval",
                "DPO / accountable owner",
                "Approve / reject / approve with conditions. Named reviewer + timestamp.",
                "Only approved DPIA can set CTRL-DPIA = present",
            ],
            [
                "9. Version",
                "System",
                "dpia_id + version. Edits create a new version; previous versions remain readable.",
                "Engine binds to a specific version hash",
            ],
            [
                "10. Review date",
                "System + DPO",
                "Next review date. Material change of processing reopens screening.",
                "Expired DPIA → CTRL-DPIA = insufficient, score does not drop",
            ],
        ],
    )

    heading(doc, "Workspace object (governance contract)", 2)
    add_para(
        doc,
        "New Pydantic models (names indicative): DpiaWorkspace, DpiaStage, DpiaRisk, DpiaControlLink. "
        "They live next to PrivacyTriage. PrivacyTriage remains the engine’s screening snapshot. "
        "The workspace is the living document.",
    )
    add_table(
        doc,
        ["Field", "Meaning"],
        [
            ["dpia_id", "Stable UUID for the DPIA work product"],
            ["assessment_id / vendor_id / tenant_id", "Lineage. One vendor may have many DPIA versions."],
            ["status", "draft | screening | in_progress | dpo_review | approved | rejected | superseded | review_due"],
            ["trigger_basis", "Copied from PrivacyTriage.dpia_decision_basis so the trigger is auditable"],
            ["screening_outcome", "required | not_required | inconclusive — DPO-owned, not checkbox-owned"],
            ["necessity_statement", "Purpose, lawful basis, proportionality, alternatives"],
            ["risks[]", "harm, data_subjects, likelihood, severity, inherent_risk"],
            ["controls[]", "control_id, description, evidence_artifact_id, residual_after_control"],
            ["dpo_residual_risk", "Very Low … Critical — DPO opinion, not engine score"],
            ["dpo_reviewer / approved_at", "Human gate. Same discipline as Jira Legal approval."],
            ["version / content_sha256 / review_by", "Immutability and recertification"],
        ],
    )

    heading(doc, "LLM role inside the DPIA Workspace", 2)
    add_para(
        doc,
        "The model may draft screening questions, risk language, and control suggestions from intake JSON "
        "and attached documents. It must not set screening_outcome, dpo_residual_risk, or approval. "
        "Unattested draft text is labelled “AI-assisted draft — not evidence.” Missing documents still "
        "return EVIDENCE NOT FOUND. This preserves ADR-2 and ADR-9.",
    )
    heading(doc, "Jira interaction", 2)
    bullet(doc, " Legal Task remains the department gate. Closing it without an approved DPIA (when indicated) is a process exception, not an engine pass.", bold_prefix="Legal ticket:")
    bullet(doc, " the Legal description links to /dpia/{dpia_id}. Reviewers work in the workspace, then close Jira.", bold_prefix="Deep link:")
    bullet(doc, " CTRL-DPIA evidence_status=present only when workspace.status=approved, version hash stored, review_by in the future, and DPO identity recorded.", bold_prefix="Control binding:")

    heading(doc, "Acceptance criteria", 2)
    bullet(doc, "A reviewer can open a DPIA from an assessment and complete all ten stages without leaving the product.")
    bullet(doc, "CTRL-DPIA does not become present from a typed reference string alone.")
    bullet(doc, "Editing an approved DPIA creates version n+1 and marks n as superseded.")
    bullet(doc, "Past review_by, or a material change to data_processed / Art. 9 / ADM, reopens screening.")
    bullet(doc, "Engine residual risk and DPO residual risk are both visible and never overwritten by each other.")

    heading(doc, "5. Gap 3 — Evidence Repository", 1)
    heading(doc, "Current state — philosophy without the archive", 2)
    add_para(
        doc,
        "The product already has the right posture: Evidence > assumptions. Missing evidence never reduces "
        "residual risk. Chat is instructed to say EVIDENCE NOT FOUND. EvidenceItem is not a Finding (ADR-8). "
        "That is the correct control-plane philosophy.",
    )
    add_para(
        doc,
        "What is still missing is the archive. EvidenceItem.document is a string (“Intake checkbox”, "
        "“SOC 2 Type II report”). CTRL-SOC2 is always missing because no file can be uploaded. "
        "A checkbox DPA stays declared_unverified. An ISO claim on a marketing page never becomes a certificate. "
        "Auditors do not accept that as a control plane. They ask: show the PDF, who reviewed it, when, and what they decided.",
    )

    heading(doc, "v2 target — artifacts bound to controls", 2)
    add_para(
        doc,
        "Introduce an Evidence Repository. Each uploaded or registered artifact is immutable (content hash). "
        "Humans bind artifacts to controls. The engine reads bindings; it does not “believe” filenames.",
    )
    add_table(
        doc,
        ["Artifact type", "Typical file", "Binds to", "Pass condition (indicative)"],
        [
            ["DPA", "Signed PDF", "CTRL-DPA", "Signed, in force, processor identity matches vendor, reviewed by Legal"],
            ["DPIA", "Workspace export or PDF", "CTRL-DPIA", "Approved version from DPIA Workspace, not an orphan Word file"],
            ["SOC 2", "Type II report PDF", "CTRL-SOC2", "In-period report; exceptions mapped; InfoSec reviewer"],
            ["ISO certificate", "ISO/IEC 27001 (and SoA if provided)", "CTRL-SOC2 / future CTRL-ISO", "Certificate in validity window; scope covers the service"],
            ["Security questionnaire", "SIG / CAIQ / custom DOCX or XLSX", "CTRL-SOC2, CTRL-MODEL, transfers", "Completed, dated, reviewer accepted; gaps become findings"],
            ["Vendor documentation", "Security whitepaper, architecture PDF", "CTRL-MODEL, CTRL-TRANSFERS", "Supports a claim; never sufficient alone for SOC 2"],
            ["Model card", "Model card / system card PDF or JSON", "CTRL-MODEL", "Identifies model, intended use, limitations, evals"],
            ["Subprocessors", "List with locations and DPAs", "CTRL-TRANSFERS", "List current; locations known; no “unknown”"],
            ["Contractual clauses", "SCCs, addenda, audit rights", "CTRL-DPA, CTRL-TRANSFERS", "Executed clauses on file; transfer tool identified"],
        ],
    )

    heading(doc, "The chain every control must have", 2)
    add_para(
        doc,
        "v1 EvidenceItem is close. v2 makes the chain explicit and complete:",
        space_after=6,
    )
    add_para(
        doc,
        "Control  →  Evidence  →  Source  →  Reviewer  →  Timestamp  →  Decision",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )
    add_table(
        doc,
        ["Link", "v1 today", "v2"],
        [
            ["Control", "control_id on EvidenceItem", "Unchanged. CTRL-* remains the vocabulary."],
            ["Evidence", "document name + excerpt string", "EvidenceArtifact (file or structured object) + excerpt + page"],
            ["Source", "Usually “intake”", "upload | intake | dpia_workspace | jira | vendor_portal | questionnaire"],
            ["Reviewer", "Missing (DecisionRecord has department approvers only)", "Named reviewer on the binding, with role"],
            ["Timestamp", "collected_at", "uploaded_at + reviewed_at (distinct)"],
            ["Decision", "evidence_status only (present / missing / …)", "pending | accepted | rejected | insufficient | not_applicable — human decision on that artifact"],
        ],
    )
    add_para(
        doc,
        "Engine rule, unchanged in spirit: only decision=accepted (and in-date) may move a control toward pass. "
        "declared_unverified remains the status for checkboxes and marketing claims. Missing evidence never "
        "reduces residual risk. A rejected artifact is not deleted; it stays in the log.",
    )

    heading(doc, "Repository objects", 2)
    add_table(
        doc,
        ["Object", "Responsibility"],
        [
            [
                "EvidenceArtifact",
                "artifact_id, tenant_id, vendor_id, type, filename, content_type, sha256, byte_size, classification, uploaded_by, uploaded_at, retention_until. Bytes in object storage (S3/GCS); DB holds metadata only.",
            ],
            [
                "EvidenceBinding",
                "The chain row: control_id, artifact_id, assessment_id, source, excerpt, page_range, reviewer, role, reviewed_at, decision, notes. This is what the engine reads.",
            ],
            [
                "EvidenceItem (API)",
                "Kept as the response projection so existing UI and evidence_pack JSON do not break. Populated from bindings, not from intake strings, once an artifact exists.",
            ],
        ],
    )

    heading(doc, "Security questionnaire as a first-class artifact", 2)
    add_para(
        doc,
        "A completed SIG / CAIQ / internal questionnaire is often the only structured vendor answer before "
        "SOC 2 is under NDA. v2 treats it as an artifact type, not as chat context. Parsed answers may "
        "pre-fill intake fields, but they do not pass controls. Each material answer still needs a binding "
        "and a reviewer decision. The questionnaire file itself is retained with hash.",
    )

    heading(doc, "Integrity and access", 2)
    bullet(doc, " SHA-256 of bytes at upload. Re-hash on download. Mismatch is a tamper event.", bold_prefix="Hash:")
    bullet(doc, " artifacts are write-once. A replacement is a new artifact_id; old bindings remain.", bold_prefix="Immutability:")
    bullet(doc, " object storage is tenant-prefixed. Signed URLs, short TTL, no public buckets.", bold_prefix="Isolation:")
    bullet(doc, " DPA / DPIA / SOC 2 are confidential. Auditors get least-privilege read.", bold_prefix="Classification:")
    bullet(doc, " follow tenant retention; support erasure of chat, not of signed legal artifacts under statutory hold.", bold_prefix="Retention:")

    heading(doc, "Acceptance criteria", 2)
    bullet(doc, "Assessor can upload a DPA PDF; Legal can accept it; CTRL-DPA becomes present with reviewer + timestamp.")
    bullet(doc, "SOC 2 remains missing until a report artifact is accepted by InfoSec. A website claim still does not pass.")
    bullet(doc, "Each control in the evidence matrix shows the six-link chain or EVIDENCE NOT FOUND.")
    bullet(doc, "evidence_pack export includes artifact hashes and binding decisions, not only narrative gaps.")
    bullet(doc, "Deleting a file in the UI is a tombstone; auditors can still see that it existed and who removed it.")

    heading(doc, "6. How the three gaps lock together", 1)
    add_para(
        doc,
        "These are not three unrelated upgrades. Persistence is the substrate. The DPIA Workspace is a "
        "governed work product that produces evidence. The repository is where that evidence, and every "
        "other artifact, becomes a control-plane object.",
    )
    add_table(
        doc,
        ["Flow", "v1", "v2"],
        [
            [
                "Personal data inferred",
                "PrivacyTriage indicated; CTRL-DPIA fails; Legal Jira opens",
                "Same trigger, plus DPIA Workspace opened at screening",
            ],
            [
                "DPA checkbox",
                "declared_unverified; score does not drop",
                "Same until a hashed PDF is accepted on CTRL-DPA",
            ],
            [
                "SOC 2",
                "Always missing",
                "Missing until Type II (or accepted equivalent) is bound and reviewed",
            ],
            [
                "Human approval",
                "Jira Done + corporate actor email",
                "Jira gates remain. DPIA approval and evidence decisions are additional, named, timestamped records",
            ],
            [
                "Audit pack",
                "JSON evidence_pack in memory",
                "Durable pack + artifact hashes + DPIA version + DecisionRecord chain",
            ],
        ],
    )

    heading(doc, "7. Suggested implementation order", 1)
    add_para(
        doc,
        "Do not start with the DPIA UI or with RAG. Without persistence, both are theatre. "
        "Effort is T-shirt size for a small senior team.",
    )
    add_table(
        doc,
        ["Step", "Deliverable", "Closes gap", "Depends on", "Effort"],
        [
            ["A", "PostgreSQL + tenant_id + RLS + append-only DecisionRecord", "1", "ADR-6 replacement", "L"],
            ["B", "Object storage + EvidenceArtifact + EvidenceBinding + upload API", "3", "A", "L"],
            ["C", "Engine reads bindings: present only if decision=accepted and in-date", "3", "B", "M"],
            ["D", "DPIA Workspace stages 1–10, versioning, DPO approval", "2", "A", "L"],
            ["E", "CTRL-DPIA binds to approved workspace version, not a reference string", "2", "C + D", "S"],
            ["F", "Security questionnaire artifact type + optional parse-to-intake", "3", "B", "M"],
            ["G", "Jira issue.key → assessment_id (replace latest-in-RAM webhook)", "1", "A", "M"],
            ["H", "SSO / RACI roles (assessor, legal, infosec, aigov, auditor)", "1", "A", "XL"],
            ["I", "PDF/DOCX evidence pack with SHA-256 of pack + included artifacts", "3", "B + D", "L"],
        ],
    )
    add_para(
        doc,
        "Steps A–E are the portfolio jump: durable, tenant-aware, DPIA as a work product, evidence as files. "
        "H is required before any real customer data. I is the Legal-facing deliverable.",
        space_before=4,
    )

    heading(doc, "8. What v2 must not do", 1)
    bullet(doc, " Let the LLM set residual risk, triage decision, DPIA approval, or evidence decision.", bold_prefix="Never:")
    bullet(doc, " Treat a filename (“soc2.pdf”) as a passed control without a reviewer decision.", bold_prefix="Never:")
    bullet(doc, " Add engine APPROVE because a DPIA workspace exists. Gates stay human.", bold_prefix="Never:")
    bullet(doc, " Lower the score because the vendor’s website claims ISO 27001.", bold_prefix="Never:")
    bullet(doc, " Use Redis or chat history as the system of record.", bold_prefix="Never:")
    bullet(doc, " Mix demo tenants with live assessments in one unscoped table.", bold_prefix="Never:")
    add_para(
        doc,
        "Preserve: Evidence → Control assessment → Risk factors → Deterministic scoring → Governance triage → Human Jira gates.",
        italic=True,
        space_before=8,
    )

    heading(doc, "9. How to talk about this in a review", 1)
    add_para(
        doc,
        "v1 is the control layer done properly under demo constraints. The scoring engine, the ADRs, and "
        "the refusal to invent evidence are the product. v2 is the operationalisation: durable tenancy, "
        "an electronic DPIA, and an evidence archive with a six-link chain on every control.",
    )
    add_para(
        doc,
        "If a reviewer asks “is this production-ready?”, the accurate answer is no. If they ask “would this "
        "design survive an audit conversation?”, the accurate answer is: the design would; the current "
        "runtime would not. That is the point of this roadmap.",
        space_before=8,
    )

    add_para(
        doc,
        "This document does not authorise skipping the deterministic engine or returning decisions to the LLM. "
        "It also does not authorise calling the current in-memory console a production-ready enterprise platform.",
        italic=True,
        color=MUTED,
        space_before=16,
    )
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
