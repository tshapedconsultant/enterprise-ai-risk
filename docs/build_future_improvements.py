"""Generate English roadmap Word doc: production store and remaining governance upgrades."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "Future_Improvements_AI_Governance.docx"
NAVY = RGBColor(0x1B, 0x21, 0x18)
ACCENT = RGBColor(0x3D, 0x5A, 0x1F)
MUTED = RGBColor(0x4A, 0x52, 0x45)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1B2118"
ROW_ALT = "F3F5EE"


def set_run(run, size=11, bold=False, color=NAVY, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, color=NAVY, italic=False, space_after=8, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"})
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=NAVY, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, color=WHITE, fill=HEADER_BG)
    for r_idx, row in enumerate(rows):
        fill = ROW_ALT if r_idx % 2 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=10, fill=fill)
    doc.add_paragraph()


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT if level == 1 else NAVY
        run.font.name = "Calibri"
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True, size=11)
        r2 = p.add_run(text)
        set_run(r2, size=11)
    else:
        r = p.add_run(text)
        set_run(r, size=11)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

    add_para(doc, "RESPONSIBLE AI GOVERNANCE  ·  ROADMAP", size=11, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Future improvements", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(
        doc,
        "From compliance assistant to Enterprise AI Governance & Risk Control Layer",
        size=13,
        italic=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16,
    )
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Date", date.today().isoformat()],
            ["Scope", "What is intentionally not in this release, and why"],
            ["Current product", "DPIA object, EvidenceItem, deterministic engine, DecisionRecord, Jira gates, evidence_pack JSON"],
            ["Out of scope here", "Durable session store (item 5) and PDF export of the pack"],
        ],
    )

    heading(doc, "1. Assessment of five governance proposals", 1)
    add_table(
        doc,
        ["#", "Proposal", "Verdict", "Status"],
        [
            ["1", "DPIA/PIA as formal object (not only has_dpa)", "Yes. GDPR is not a checkbox.", "Done (API v1.1)"],
            ["2", "EvidenceItem separate from Findings", "Yes. Finding ≠ proof.", "Done (evidence_items)"],
            ["3", "LLM must not set risk/decision", "Yes. Without this it is a compliance chatbot.", "Done (app/scoring.py)"],
            ["4", "DecisionRecord / audit trail", "Yes. Enterprise audit and sales.", "Done + Jira human gates (v1.5)"],
            ["5", "Replace in-memory store", "Yes for production. Not for demo.", "Deferred — sessions keyed by assessment_id"],
        ],
    )
    add_para(
        doc,
        "The evidence pack (assessment, matrix, GDPR/DPIA, AI Act, NIST, ISO, findings, remediation, Jira, audit trail) "
        "is already generated as JSON (evidence_pack). Missing: immutable PDF and durable storage.",
    )

    heading(doc, "2. Item 5 — session store (production)", 1)
    add_para(
        doc,
        "Today app/store.py holds sessions keyed by assessment_id (TTL-capped). Fine for demo and interviews. "
        "Not fine for multiple workers, restarts, multiple users, or audits asking who decided what on Tuesday.",
    )
    heading(doc, "Target design", 2)
    bullet(doc, " assessments, intake, evidence_items, decision_record, evidence_pack. Stable relational IDs.", bold_prefix="PostgreSQL:")
    bullet(doc, " UI session, short locks, cache of latest pack by assessment_id. Not source of truth.", bold_prefix="Redis:")
    bullet(doc, " each DecisionRecord written once, never updated. Correction = new record with supersedes_id.", bold_prefix="Immutable audit log:")
    bullet(doc, " alternative: append-only JSONL or DB triggers blocking UPDATE/DELETE.", bold_prefix="Event store:")
    heading(doc, "Must persist", 2)
    add_table(
        doc,
        ["Object", "Why RAM is not enough"],
        [
            ["VendorInput + version", "Reproduce decision: same facts, same engine, same result."],
            ["EvidenceItem", "Auditor asks for proof, not chat prose."],
            ["DecisionRecord", "Who, when, score, rules, engine version, human approvers."],
            ["GovernanceEvidencePack", "Legal deliverable; hash for integrity."],
            ["Chat history (optional)", "Lost on F5 except summary. If stored in prod: retention + minimization."],
        ],
    )
    heading(doc, "Acceptance criteria", 2)
    bullet(doc, "Restarting uvicorn does not erase assessments.")
    bullet(doc, "Two workers serve the same assessment_id.")
    bullet(doc, "Direct UPDATE on DecisionRecord in DB must fail or emit tamper event.")
    bullet(doc, "Export pack with SHA-256 hash and engine version (deterministic-rules-v1, v2, …).")

    heading(doc, "3. Evidence pack — remaining work (PDF / portal)", 1)
    add_para(doc, "JSON already has the sellable tree. Operational differentiation is an artifact Legal can archive without the console:")
    bullet(doc, "PDF or DOCX on assessment close (corporate template, pagination, signatures).")
    bullet(doc, "Evidence portal with NDA (vendor SOC 2 attached, not only missing).")
    bullet(doc, "Jira: Epic + sub-tasks payload, optional POST (JIRA_BASE_URL + token), inbound webhook /api/v1/webhooks/jira. No credentials = dry-run.")
    bullet(doc, "EU AI Act Annex III classification as structured object, not a paragraph.")

    heading(doc, "4. Other production improvements", 1)
    add_table(
        doc,
        ["Topic", "Why", "Priority", "Effort"],
        [
            ["Corporate SSO and RACI roles", "No owner means no control layer.", "High", "L"],
            ["Attachments (DPA/DPIA PDF) with hash", "Move from declared_unverified to present.", "High", "L"],
            ["RAG on approved internal KB", "OSINT and contracts; never model memory.", "Medium", "XL"],
            ["Multitenancy / environments", "Do not mix demos with live cases.", "High", "L"],
            ["Chat retention and erasure", "Chat is also processing.", "Medium", "M"],
            ["Engine calibration on historical cases", "Rules v1 are conservative by design.", "Medium", "M"],
        ],
    )

    heading(doc, "5. Commercial framing", 1)
    add_para(doc, "Do not position as:", bold=True, space_after=4)
    add_para(doc, "“AI Compliance Assistant” / chatbot that scores vendors.", italic=True)
    add_para(doc, "Position as:", bold=True, space_after=4, space_before=8)
    add_para(
        doc,
        "Enterprise AI Governance & Risk Control Layer: facts → controls → evidence → deterministic score → "
        "traceable triage → remediation (Jira) → audit pack.",
        italic=True,
    )
    add_para(
        doc,
        "Public “RAG + scoring” products already exist. The difference is operational governance: "
        "the model explains, the engine triages, the log is durable, Legal gets a pack.",
        space_before=8,
    )

    heading(doc, "6. Suggested implementation order", 1)
    add_table(
        doc,
        ["Step", "Deliverable", "Depends on", "Effort"],
        [
            ["A", "PostgreSQL + assessment_id + append-only DecisionRecord", "Item 5", "L"],
            ["B", "Attachments and evidence_status present", "A", "L"],
            ["C", "Harden Jira automation (accountId mapping)", "A", "M"],
            ["D", "PDF Evidence Pack with hash", "A + Legal template", "L"],
            ["E", "SSO and environments", "A", "XL"],
            ["F", "Redis as UI session cache only", "A (never source of truth)", "S"],
        ],
    )

    add_para(
        doc,
        "Effort uses T-shirt sizes (S / M / L / XL) for a small senior team. XL includes organisational work, not only code.",
        italic=True,
        color=MUTED,
        space_before=4,
    )

    add_para(
        doc,
        "This document does not authorize skipping the deterministic engine or returning decisions to the LLM. "
        "All future work must preserve: Evidence → Control assessment → Risk factors → Deterministic scoring → Governance triage → Human Jira gates.",
        italic=True,
        space_before=12,
        color=MUTED,
    )
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
