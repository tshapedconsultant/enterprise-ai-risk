"""Generate English user guide as Word document (v1.7 — tabbed report, audit chips, disclosure rows)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "User_Guide_AI_Vendor_Risk.docx"
NAVY = RGBColor(0x1B, 0x21, 0x18)
ACCENT = RGBColor(0x3D, 0x5A, 0x1F)
MUTED = RGBColor(0x4A, 0x52, 0x45)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1B2118"
ROW_ALT = "F3F5EE"


def set_run(run, size=11, bold=False, color=NAVY, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, color=NAVY, italic=False, space_after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"})
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=NAVY, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, color=WHITE, fill=HEADER_BG)
    for r_idx, row in enumerate(rows):
        fill = ROW_ALT if r_idx % 2 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=10, fill=fill)
    doc.add_paragraph()


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT if level == 1 else NAVY
        run.font.name = "Calibri"


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, size=11)


def build():
    doc = Document()
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], margin, Cm(2.2))

    add_para(doc, "RESPONSIBLE AI GOVERNANCE", size=11, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "User Guide — AI Vendor Risk Console", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(
        doc,
        "For business, legal, compliance, and security readers",
        size=12,
        italic=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16,
    )
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Version", "1.7"],
            ["Date", date.today().isoformat()],
            ["Audience", "Business, Legal, Compliance, Security, Leadership"],
            ["Purpose", "How to use the console, what it decides, what it does not decide"],
        ],
    )

    heading(doc, "1. One-page summary", 1)
    add_para(
        doc,
        "The Enterprise AI Vendor Risk Console is a governance control layer for third-party AI vendors. "
        "You enter facts about a proposed integration; the rules engine triages risk and opens Jira work for "
        "Legal, SecOps, and AI Governance. The chat explains the report but cannot approve a vendor.",
    )
    bullet(doc, "Engine triage only: PENDING REVIEW, REQUIRES REMEDIATION, or ESCALATE — never automatic APPROVE.")
    bullet(doc, "Humans approve in Jira (@adevinta.com). When all department Tasks close, workflow becomes DEPARTMENT_GATES_COMPLETED (awaiting final approval).")
    bullet(doc, "Without OPENAI_API_KEY, simulator mode uses business rules — suitable for demos.")

    heading(doc, "2. Console layout (v1.7)", 1)
    add_para(
        doc,
        "The window is two columns. The right column is the only scroll area: no nested scrollbars inside "
        "findings or evidence panels. The decision bar stays pinned at the top; the chat composer stays pinned at the bottom.",
    )
    add_table(
        doc,
        ["Area", "What you see"],
        [
            ["Left panel — intake", "Vendor form with high-contrast fields. Required fields first; Privacy/DPIA and additional info in expandable sections."],
            ["Decision bar (pinned)", "Triage decision, residual risk, vendor name, Copy report button."],
            ["Audit trail", "Labelled grid: triage, workflow, engine score, residual risk, engine version, scoring method."],
            ["Gate chips", "Legal, SecOps, AI Governance — amber Pending, green with approver email, red No ticket."],
            ["Report tabs", "Findings | Missing evidence | Jira orchestration. Count badge on each tab. One panel visible at a time."],
            ["Expandable rows", "Headline in bold (e.g. CTRL-DPIA + missing chip). Click the row to read the full rationale, document, and source."],
            ["Risk assistant card", "Structured summary after assess: decision grid, count chips, gate status. Not a plain-text log."],
            ["Chat", "Follow-up questions below the card. Suggestion chips above the composer."],
        ],
    )
    heading(doc, "Reading rules", 2)
    bullet(doc, "Tabs, not columns — each report section uses the full width.")
    bullet(doc, "Headline first — long explanations stay behind an expandable row (progressive disclosure).")
    bullet(doc, "Chips carry status — colour is always paired with a word (Pending, approved, missing).")
    bullet(doc, "Copy report — plain text for email/tickets. The on-screen card is visual only.")

    heading(doc, "3. Step-by-step", 1)
    bullet(doc, "Open http://127.0.0.1:8000 (or your deployed URL).")
    bullet(doc, "Replace the example vendor data with your real case.")
    bullet(doc, "Check DPA only if a signed agreement exists with your entity.")
    bullet(doc, "Expand Privacy and DPIA when personal data, transfers, or automated decisions apply.")
    bullet(doc, "Click Assess vendor.")
    bullet(doc, "Read the decision bar and audit trail chips.")
    bullet(doc, "Open each report tab. Expand rows where you need the full rationale.")
    bullet(doc, "Use suggestion chips or chat for follow-up questions.")
    bullet(doc, "Copy report for email or tickets if needed.")

    heading(doc, "4. Triage decisions", 1)
    add_table(
        doc,
        ["Code", "Meaning"],
        [
            ["PENDING REVIEW", "No critical red flag; departments must still close Jira gates"],
            ["REQUIRES REMEDIATION", "Material gaps (DPA, DPIA, score, etc.)"],
            ["ESCALATE TO AI GOVERNANCE / LEGAL / SECURITY", "Personal data without DPA, Art. 9, or score 5"],
            ["DEPARTMENT_GATES_COMPLETED", "All three department Tasks closed — awaiting final decision-maker"],
            ["APPROVE WITH CONDITIONS", "Human only — reserved for a later final decision-maker"],
        ],
    )

    heading(doc, "5. Jira orchestration", 1)
    add_para(doc, "Each assessment proposes:")
    bullet(doc, "Parent Epic — holds triage context (not an approval by itself).")
    bullet(doc, "Legal Task — DPA and DPIA (label legal-review, dpo@adevinta.com).")
    bullet(doc, "SecOps Task — SOC 2 Type II (label infosec, secops@adevinta.com).")
    bullet(doc, "AI Governance Task — bias and explainability (label ai-governance-review).")
    add_para(
        doc,
        "With Jira credentials configured, tickets are created via REST API. Otherwise payloads are dry-run in the UI. "
        "Closing a Task in Jira should trigger POST /api/v1/webhooks/jira with a corporate approver email.",
        italic=True,
        color=MUTED,
    )

    heading(doc, "6. FAQ", 1)
    add_para(doc, "What if the vendor has no SOC 2 but has ISO 27001?", bold=True, space_after=4)
    add_para(
        doc,
        "A website claim is not evidence. Keep the SecOps Jira Task open. SecOps reviews the actual ISO 27001 "
        "certificate and latest audit. If they accept that pack as equivalent independent assurance, they close the "
        "infosec ticket. The engine score does not drop until a file is attached in a later release. SOC 2 Type II "
        "and ISO 27001 are not the same artefact; SecOps decides equivalence.",
    )
    add_para(doc, "Who closes the AI Governance ticket if we have no dedicated AI Governance team?", bold=True, space_after=4)
    add_para(
        doc,
        "Do not skip the gate. Assign ai-governance-review to the closest named RACI owner (CISO, DPO plus business "
        "owner, or a risk committee member) with an @adevinta.com account. Closing it means someone accepted model risk.",
    )
    add_para(doc, "Why is residual risk still High after I checked the DPA box?", bold=True, space_after=4)
    add_para(
        doc,
        "A checkbox is a declaration. Without the signed PDF, the control is declared_unverified. Missing evidence never lowers the score.",
    )
    add_para(doc, "Can chat or the LLM approve the vendor?", bold=True, space_after=4)
    add_para(
        doc,
        "No. Chat explains the report. The structured assistant card and the decision bar show engine triage only. "
        "Closing Legal, SecOps, and AI Governance Tasks produces DEPARTMENT_GATES_COMPLETED, which still awaits a final decision-maker.",
    )

    heading(doc, "7. Legal disclaimer", 1)
    add_para(
        doc,
        "This console does not provide legal advice. It does not replace a DPO opinion, signed contracts, "
        "or independent security audits. Residual risk does not decrease because a vendor claims compliance on a website.",
        italic=True,
    )

    heading(doc, "8. Document history", 1)
    add_table(
        doc,
        ["Version", "Date", "Changes"],
        [
            ["1.7", date.today().isoformat(), "Tabbed report, audit chips, disclosure rows, assistant card, form contrast"],
            ["1.6", "2026-08-18", "Console map, FAQ (ISO 27001 vs SOC 2, missing AI Gov team)"],
            ["1.5", "2026-08-14", "English UI; Jira Epic + Tasks; webhook approvals; comprehensive docs"],
        ],
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
