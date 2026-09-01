"""Render Markdown documentation to DOCX from one canonical source.

Usage:
    python scripts/export_docx.py docs/USER_GUIDE.md
    python scripts/export_docx.py docs/USER_GUIDE.md docs/ARCHITECTURE.md --out-dir artifacts/docs

Markdown remains the only maintained source. Generated .docx files are build
artifacts and are intentionally ignored by Git.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _add_inline(paragraph, text: str) -> None:
    """Add a small, deterministic subset of Markdown inline formatting."""
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            label, url = token[1:].split("](", 1)
            paragraph.add_run(f"{label} ({url[:-1]})")
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def render_markdown(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    document.core_properties.title = source.stem.replace("_", " ").title()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    in_code = False
    index = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            index += 1
            continue
        if in_code:
            run = document.add_paragraph().add_run(raw)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            index += 1
            continue
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = _table_rows(lines, index)
            if rows:
                width = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for col_index, value in enumerate(row):
                        cell = table.cell(row_index, col_index)
                        _add_inline(cell.paragraphs[0], value)
                        if row_index == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            document.add_heading(heading.group(2), level=min(len(heading.group(1)), 3))
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            paragraph = document.add_paragraph(
                style="List Bullet" if bullet else "List Number"
            )
            _add_inline(paragraph, (bullet or numbered).group(1))
            index += 1
            continue
        paragraph = document.add_paragraph()
        _add_inline(paragraph, stripped)
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def export(sources: Iterable[Path], out_dir: Path | None) -> list[Path]:
    outputs = []
    for source in sources:
        if not source.is_file():
            raise FileNotFoundError(source)
        output = (out_dir or source.parent) / f"{source.stem}.docx"
        render_markdown(source, output)
        outputs.append(output)
    return outputs


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("sources", nargs="+", type=Path)
    parser.add_argument("--out-dir", type=Path)
    args = parser.parse_args()
    for output in export(args.sources, args.out_dir):
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
